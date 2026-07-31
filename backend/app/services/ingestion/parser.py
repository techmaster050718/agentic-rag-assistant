from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.document_loaders import BaseLoader
from langchain_core.documents import Document
from docx import Document as DocxDocument
from app.core.logging import get_logger

logger = get_logger(__name__)

class DocxLoader(BaseLoader):
    def __init__(self, file_path: str):
        self.file_path = file_path
    def load(self) -> list[Document]:
        doc = DocxDocument(self.file_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return [Document(page_content=text, metadata={"source": self.file_path, "file_type": "docx"})]

def parse_document(file_path: str, file_type: str) -> list[Document]:
    try:
        if file_type == "pdf": loader = PyPDFLoader(file_path)
        elif file_type == "txt": loader = TextLoader(file_path, encoding="utf-8")
        elif file_type in ["doc", "docx"]: loader = DocxLoader(file_path)
        else: raise ValueError(f"Unsupported file type: {file_type}")
        documents = loader.load()
        logger.info("Parsed document", file_path=file_path, num_pages=len(documents))
        return documents
    except Exception as e:
        logger.error("Failed to parse document", file_path=file_path, error=str(e))
        raise
